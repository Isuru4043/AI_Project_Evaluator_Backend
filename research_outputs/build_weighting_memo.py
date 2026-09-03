from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).with_name('VivaSense_50-35-15_Weighting_Research_Justification.docx')

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
GRAY = RGBColor(90, 101, 115)
LIGHT = 'F2F4F7'
PALE_BLUE = 'E8EEF5'
WHITE = RGBColor(255, 255, 255)


def font(run, size=11, bold=False, italic=False, color=NAVY):
    run.font.name = 'Calibri'
    run._element.get_or_add_rPr().rFonts.set(qn('w:ascii'), 'Calibri')
    run._element.get_or_add_rPr().rFonts.set(qn('w:hAnsi'), 'Calibri')
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        p_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement('w:tblHeader')
    header.set(qn('w:val'), 'true')
    tr_pr.append(header)


def margins(cell, top=100, start=130, bottom=100, end=130):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for tag, value in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tc_mar.find(qn(f'w:{tag}'))
        if node is None:
            node = OxmlElement(f'w:{tag}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in('w:tblW')
    if tbl_w is None:
        tbl_w = OxmlElement('w:tblW')
        tbl_pr.append(tbl_w)
    tbl_w.set(qn('w:w'), str(sum(widths_dxa)))
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_ind = tbl_pr.first_child_found_in('w:tblInd')
    if tbl_ind is None:
        tbl_ind = OxmlElement('w:tblInd')
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn('w:w'), '120')
    tbl_ind.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement('w:gridCol')
        col.set(qn('w:w'), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in('w:tcW')
            if tc_w is None:
                tc_w = OxmlElement('w:tcW')
                tc_pr.append(tc_w)
            tc_w.set(qn('w:w'), str(widths_dxa[idx]))
            tc_w.set(qn('w:type'), 'dxa')
            margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_link(paragraph, label, url):
    part = paragraph.part
    rel_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), rel_id)
    run = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '2E74B5')
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rpr.extend([color, underline])
    text = OxmlElement('w:t')
    text.text = label
    run.extend([rpr, text])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run('Page ')
    font(run, 9, color=GRAY)
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    paragraph._p.append(fld)


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        font(r1, bold=True)
        r2 = p.add_run(text[len(bold_lead):])
        font(r2)
    else:
        r = p.add_run(text)
        font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.167
    font(p.add_run(text))


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = section.bottom_margin = Inches(1)
section.left_margin = section.right_margin = Inches(1)
section.header_distance = section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = NAVY
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for name, size, color, before, after in [
    ('Heading 1', 16, BLUE, 16, 8),
    ('Heading 2', 13, BLUE, 12, 6),
    ('Heading 3', 12, RGBColor(31, 77, 120), 8, 4),
]:
    st = styles[name]
    st.font.name = 'Calibri'
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = color
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
font(header.add_run('VivaSense | Assessment Design Research Memo'), 9, bold=True, color=GRAY)
footer = section.footer.paragraphs[0]
add_page_field(footer)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(18)
p.paragraph_format.space_after = Pt(4)
font(p.add_run('RESEARCH MEMO'), 10, bold=True, color=BLUE)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(5)
font(p.add_run('Justifying the 50/35/15 Answer-Scoring Weights'), 24, bold=True, color=NAVY)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(14)
font(p.add_run('Evidence base, mathematical consequences, limitations, and evaluator-ready defence'), 12, color=GRAY)
for label, value in [('System', 'VivaSense AI Project Evaluator and Viva Examiner'), ('Formula', 'S_a = 0.50C + 0.35D + 0.15K'), ('Dimensions', 'Correctness (C), Depth (D), Consistency (K)'), ('Prepared', '2 September 2026')]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    font(p.add_run(label + ': '), bold=True)
    font(p.add_run(value))

doc.add_heading('Executive conclusion', level=1)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.12)
p.paragraph_format.right_indent = Inches(0.12)
p.paragraph_format.space_before = Pt(3)
p.paragraph_format.space_after = Pt(8)
shade_paragraph(p, PALE_BLUE)
font(p.add_run('Defensible conclusion. '), bold=True, color=NAVY)
font(p.add_run('No credible source mandates the exact 50/35/15 vector for this rubric. The defensible claim is that research supports explicit analytic weighting, the selected priority order, and empirical validation; the exact percentages are transparent expert-informed design weights with deliberate score ceilings.'), color=NAVY)

doc.add_heading('What the research directly supports', level=1)
add_body(doc, 'Explicit weights are preferable to hidden combination rules. Tomas et al. found that implicit holistic marking could diverge from intended learning outcomes, while published analytic criteria and weights improve transparency.')
add_body(doc, 'The same study produced a useful external comparator: data-derived criterion importance was approximately 51-52% for knowledge and understanding, 31-32% for critical thinking, and 17% for lower-order writing criteria. These are not identical constructs, but the proportional hierarchy closely resembles VivaSense\'s 50/35/15 design.')
add_body(doc, 'Wainer and Thissen argue that component weights affect reliability and validity and should reflect both logical validity and empirical reliability. Graves similarly connects fair weighting with measurable alignment to learning outcomes.')
add_body(doc, 'Revised Bloom\'s taxonomy distinguishes factual knowledge from increasingly complex cognitive processes, supporting separate correctness and depth dimensions. Computing-education research has also validated multidimensional scoring for code explanations rather than relying on one holistic score.')
add_body(doc, 'Structured viva research supports consistent scoring rules: a systematic review of 24 studies reported reliability around 0.75-0.80 in two structured-viva settings, compared with 0.50 for traditional viva. Coursework-authentication research supports using oral answers to corroborate authorship of submitted work.')

doc.add_heading('Why each percentage is intentional', level=1)
doc.add_heading('50% correctness: primary direct evidence', level=2)
add_body(doc, 'Correctness answers the fundamental question: is the technical claim true and supported by the report, code, and rubric? Giving it half the score makes factual validity dominant while leaving room to distinguish levels of understanding. It also closely tracks the 51-52% knowledge-and-understanding importance observed by Tomas et al.; this is corroboration rather than an identical construct mapping.')
doc.add_heading('35% depth: substantial evidence of understanding', level=2)
add_body(doc, 'Depth captures mechanisms, justification, trade-offs, application, and limitations. Its substantial weight reflects the viva\'s purpose of distinguishing understanding from recall, while keeping it below correctness because elaborate reasoning cannot repair a false technical answer.')
doc.add_heading('15% consistency: bounded corroborating evidence', level=2)
add_body(doc, 'Consistency checks coherence with the submission and earlier answers, supporting authenticity and contradiction detection. It is indirect and noisier than correctness or depth because speech recognition, question ambiguity, changed assumptions, and limited transcript history can create apparent inconsistency. Its influence is therefore deliberately capped.')

doc.add_heading('The mathematical safeguards', level=1)
table = doc.add_table(rows=1, cols=3)
set_table_geometry(table, [2400, 2280, 4680])
headers = ['Missing dimension', 'Maximum score', 'Designed interpretation']
for i, text in enumerate(headers):
    shade(table.cell(0, i), LIGHT)
    p = table.cell(0, i).paragraphs[0]
    font(p.add_run(text), 10, bold=True)
mark_header_row(table.rows[0])
rows = [
    ('Correctness = 0', '50%', 'A detailed, consistent but false answer cannot receive a high mark.'),
    ('Depth = 0', '65%', 'A correct but superficial or memorized answer cannot receive an excellent mark.'),
    ('Consistency = 0', '85%', 'A noisy corroborating signal penalizes but cannot erase strong direct evidence.'),
]
for vals in rows:
    cells = table.add_row().cells
    for i, text in enumerate(vals):
        p = cells[i].paragraphs[0]
        font(p.add_run(text), 10)
        if i == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_table_geometry(table, [2400, 2280, 4680])

doc.add_heading('Transparent derivation', level=1)
add_bullet(doc, 'Stage 1: allocate 50% to the primary direct evidence, correctness.')
add_bullet(doc, 'Stage 2: divide the remaining 50% in a 70:30 ratio between direct evidence of understanding and corroborating evidence.')
add_bullet(doc, 'The result is 35% for depth and 15% for consistency, preserving C > D > K.')
add_body(doc, 'This is a constrained expert-judgement rule. Research recognizes expert-derived percentage weights as an acceptable initial approach, but also shows why those weights should later be checked against actual examiner judgements.')

doc.add_heading('Evaluator-ready defence', level=1)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.12)
p.paragraph_format.right_indent = Inches(0.12)
p.paragraph_format.space_before = Pt(3)
p.paragraph_format.space_after = Pt(8)
shade_paragraph(p, 'F4F6F9')
defence = ('The weights were defined as an explicit analytic combination rule aligned with the purpose of a technical viva. Correctness receives 50% because it is the most direct and necessary evidence of technical competence. Depth receives 35% because educational taxonomies distinguish factual recall from explanation, application, analysis, and evaluation. Consistency receives 15% because coherence with the report, code, and previous answers strengthens authenticity, but it is indirect and more vulnerable to transcription and context errors. The score ceilings are deliberate: zero correctness limits the score to 50%, zero depth limits it to 65%, while inconsistency alone cannot reduce an otherwise correct and deep answer below 85%. Research supports expert-derived percentage weights followed by empirical checking, and an authentic marking study independently found a similar hierarchy of approximately 51-52% knowledge and understanding, 31-32% critical thinking, and 17% secondary criteria. We therefore describe 50/35/15 as a transparent, literature-informed initial weighting to be calibrated against independent examiner judgements, not as a universal constant.')
font(p.add_run(defence), 10.5, italic=True)

doc.add_heading('Required limitation and validation', level=1)
add_body(doc, 'Do not claim that a publication prescribed the exact vector. For a final empirical claim, have at least two qualified examiners score a frozen answer set; compare 50/35/15 with equal and fitted weights on held-out cases; report inter-rater agreement, correlation, mean absolute error, and sensitivity to small weight changes. If correctness must be non-compensable, evaluate a minimum-correctness gate in addition to the weighted sum.')

doc.add_heading('References', level=1)
refs = [
    ('Abuzied & Nabag (2023), structured-viva systematic review and meta-analysis', 'https://doi.org/10.1186/s12909-023-04524-6'),
    ('Cao & Zahid (2025), automated viva for coursework authentication', 'https://kclpure.kcl.ac.uk/portal/en/publications/automated-viva-voce-using-generative-ai-for-student-coursework-au/'),
    ('Chen et al. (2020), validated code-explanation scoring rubric', 'https://doi.org/10.1145/3328778.3366879'),
    ('Graves (2026), optimal assessment weighting and learning-outcome alignment', 'https://doi.org/10.14434/josotl.v26i1.37647'),
    ('Krathwohl (2002), overview of revised Bloom taxonomy', 'https://doi.org/10.1207/s15430421tip4104_2'),
    ('Mislevy & Haertel (2006), Evidence-Centered Design', 'https://doi.org/10.1111/j.1745-3992.2006.00075.x'),
    ('Tomas et al. (2019), modeling holistic marks with analytic rubrics', 'https://doi.org/10.3389/feduc.2019.00089'),
    ('Wainer & Thissen (2004), weighting component assessments', 'https://pubmed.ncbi.nlm.nih.gov/15446296/'),
]
for label, url in refs:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(5)
    add_link(p, label, url)

doc.core_properties.title = 'VivaSense 50/35/15 Weighting Research Justification'
doc.core_properties.subject = 'Research-informed justification for answer-scoring weights'
doc.core_properties.author = 'VivaSense Project Team'
doc.save(OUT)
print(OUT)
