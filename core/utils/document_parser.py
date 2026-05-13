import os
import fitz  # PyMuPDF
from docx import Document


def extract_text_from_file(file_path: str) -> str:
    """
    Extracts plain text from a PDF or DOCX file.

    Args:
        file_path: Absolute path to the file on disk.

    Returns:
        Extracted text as a single string.

    Raises:
        ValueError: If the file type is not supported.
        FileNotFoundError: If the file does not exist.
    """

    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.pdf':
        return _extract_from_pdf(file_path)
    elif ext in ('.docx', '.doc'):
        return _extract_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Only PDF and DOCX are supported.")


def _extract_from_pdf(file_path: str) -> str:
    """Extract text from a PDF using PyMuPDF."""
    text_parts = []

    with fitz.open(file_path) as doc:
        for page in doc:
            text = page.get_text("text")
            if text.strip():
                text_parts.append(text.strip())

    return "\n\n".join(text_parts)


def _extract_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    doc = Document(file_path)
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())

    return "\n\n".join(text_parts)


def extract_text_from_bytes(file_content: bytes, filename: str) -> str:
    """
    Extracts plain text from file content bytes (PDF or DOCX).

    Works with cloud storage backends (e.g., Azure Blob Storage)
    where files don't have a local filesystem path.

    Args:
        file_content: Raw bytes of the file.
        filename: Original filename (used to determine file type).

    Returns:
        Extracted text as a single string.

    Raises:
        ValueError: If the file type is not supported.
    """
    import io

    ext = os.path.splitext(filename)[1].lower()

    if ext == '.pdf':
        return _extract_from_pdf_bytes(file_content)
    elif ext in ('.docx', '.doc'):
        return _extract_from_docx_bytes(file_content)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Only PDF and DOCX are supported.")


def _extract_from_pdf_bytes(file_content: bytes) -> str:
    """Extract text from PDF bytes using PyMuPDF."""
    import io
    text_parts = []

    with fitz.open(stream=file_content, filetype="pdf") as doc:
        for page in doc:
            text = page.get_text("text")
            if text.strip():
                text_parts.append(text.strip())

    return "\n\n".join(text_parts)


def _extract_from_docx_bytes(file_content: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    import io
    doc = Document(io.BytesIO(file_content))
    text_parts = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())

    return "\n\n".join(text_parts)


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    """
    Splits extracted text into overlapping chunks for embedding.

    Each chunk is roughly chunk_size characters, with overlap characters
    shared between consecutive chunks. This ensures context isn't lost
    at chunk boundaries when searching with FAISS.

    Args:
        text: The full extracted text.
        chunk_size: Target character length per chunk.
        overlap: Number of characters to overlap between chunks.

    Returns:
        List of text chunk strings.
    """
    if not text.strip():
        return []

    chunks = []
    start = 0

    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]

        if chunk.strip():
            chunks.append(chunk.strip())

        start += chunk_size - overlap

    return chunks